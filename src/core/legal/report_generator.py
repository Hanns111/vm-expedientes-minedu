"""
Generador automático de reportes normativos en PDF y Word
Usando Jinja2 + python-docx + WeasyPrint (100% gratuito)
"""
import logging
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass
from enum import Enum
import json

# Templates
from jinja2 import Environment, FileSystemLoader, Template

# Word generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# PDF generation
try:
    import weasyprint
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

logger = logging.getLogger(__name__)

class ReportFormat(Enum):
    """Formatos de reporte disponibles"""
    WORD = "word"
    PDF = "pdf"
    HTML = "html"

@dataclass
class ReportTemplate:
    """Template de reporte normativo"""
    name: str
    title: str
    description: str
    sections: List[str]
    format_type: ReportFormat
    template_file: str

class NormativeReportGenerator:
    """
    Generador automático de reportes normativos profesionales
    Con estructura formal y trazabilidad completa
    """
    
    def __init__(self, templates_path: Optional[Path] = None, output_path: Optional[Path] = None):
        self.templates_path = templates_path or Path(__file__).parent / "templates"
        self.output_path = output_path or Path(__file__).parent / "reports"
        
        # Crear directorios si no existen
        self.templates_path.mkdir(exist_ok=True)
        self.output_path.mkdir(exist_ok=True)
        
        # Configurar Jinja2
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_path)),
            autoescape=True
        )
        
        # Templates disponibles
        self.available_templates = self._load_available_templates()
        
        # Crear templates básicos si no existen
        self._ensure_basic_templates()
        
        logger.info(f"📄 NormativeReportGenerator inicializado - Templates: {len(self.available_templates)}")
    
    def generate_legal_report(
        self,
        query: str,
        analysis_result: Dict[str, Any],
        template_name: str = "legal_analysis",
        format_type: ReportFormat = ReportFormat.WORD,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Path:
        """
        Generar reporte legal completo
        
        Args:
            query: Consulta original del usuario
            analysis_result: Resultado del análisis legal
            template_name: Nombre del template a usar
            format_type: Formato de salida (WORD, PDF, HTML)
            metadata: Metadatos adicionales
        
        Returns:
            Path al archivo generado
        """
        try:
            # Preparar datos para el template
            report_data = self._prepare_report_data(query, analysis_result, metadata)
            
            # Generar según formato
            if format_type == ReportFormat.WORD and DOCX_AVAILABLE:
                return self._generate_word_report(report_data, template_name)
            elif format_type == ReportFormat.PDF and PDF_AVAILABLE:
                return self._generate_pdf_report(report_data, template_name)
            else:
                return self._generate_html_report(report_data, template_name)
                
        except Exception as e:
            logger.error(f"❌ Error generando reporte: {e}")
            raise
    
    def _prepare_report_data(
        self, 
        query: str, 
        analysis_result: Dict[str, Any], 
        metadata: Optional[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """Preparar datos estructurados para el reporte"""
        
        current_time = datetime.now()
        
        report_data = {
            # Información básica
            "query": query,
            "generation_date": current_time.strftime("%d/%m/%Y"),
            "generation_time": current_time.strftime("%H:%M:%S"),
            "trace_id": analysis_result.get("trace_id", f"report_{int(current_time.timestamp())}"),
            
            # Análisis legal
            "response": analysis_result.get("response", ""),
            "sources": analysis_result.get("sources", []),
            "confidence": analysis_result.get("confidence", 0.0),
            "documents_found": analysis_result.get("documents_found", 0),
            "method": analysis_result.get("method", ""),
            "agent_used": analysis_result.get("agent_used", ""),
            
            # Estructura legal
            "legal_basis": self._extract_legal_basis(analysis_result),
            "normative_references": self._extract_normative_references(analysis_result),
            "conclusions": self._extract_conclusions(analysis_result),
            
            # Metadatos
            "metadata": metadata or {},
            "user_info": metadata.get("user", {}) if metadata else {},
            "institution": "Ministerio de Educación del Perú",
            "system_version": "RAG MINEDU v1.4.0"
        }
        
        return report_data
    
    def _extract_legal_basis(self, analysis_result: Dict[str, Any]) -> List[Dict[str, str]]:
        """Extraer base legal de las fuentes"""
        legal_basis = []
        sources = analysis_result.get("sources", [])
        
        for source in sources:
            if isinstance(source, dict):
                legal_item = {
                    "norm": source.get("source", ""),
                    "article": source.get("title", ""),
                    "content": source.get("content", "")[:200] + "...",
                    "relevance": source.get("score", 0.0)
                }
                legal_basis.append(legal_item)
        
        return legal_basis
    
    def _extract_normative_references(self, analysis_result: Dict[str, Any]) -> List[str]:
        """Extraer referencias normativas específicas"""
        response = analysis_result.get("response", "")
        
        # Buscar patrones de referencias normativas
        import re
        patterns = [
            r"Decreto Supremo N°?\s*\d+[-\d]*[-\w]*",
            r"Ley N°?\s*\d+",
            r"Directiva N°?\s*\d+[-\d]*[-\w]*",
            r"Artículo \d+",
            r"Resolución Ministerial N°?\s*\d+[-\d]*[-\w]*"
        ]
        
        references = []
        for pattern in patterns:
            matches = re.findall(pattern, response, re.IGNORECASE)
            references.extend(matches)
        
        return list(set(references))  # Eliminar duplicados
    
    def _extract_conclusions(self, analysis_result: Dict[str, Any]) -> List[str]:
        """Extraer conclusiones del análisis"""
        response = analysis_result.get("response", "")
        
        # Buscar secciones de conclusiones o respuestas directas
        conclusions = []
        
        # Dividir respuesta en párrafos
        paragraphs = response.split('\n')
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if paragraph and len(paragraph) > 50:  # Párrafos sustanciales
                # Identificar conclusiones
                if any(keyword in paragraph.lower() for keyword in 
                       ['por tanto', 'en conclusión', 'se establece', 'corresponde', 'debe']):
                    conclusions.append(paragraph)
        
        # Si no hay conclusiones explícitas, usar párrafos finales
        if not conclusions and paragraphs:
            conclusions = [p.strip() for p in paragraphs[-2:] if p.strip() and len(p.strip()) > 30]
        
        return conclusions
    
    def _generate_word_report(self, report_data: Dict[str, Any], template_name: str) -> Path:
        """Generar reporte en formato Word"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx no está disponible")
        
        # Crear documento Word
        doc = Document()
        
        # Configurar estilos
        self._setup_word_styles(doc)
        
        # Encabezado
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"{report_data['institution']} - Reporte Legal"
        
        # Título principal
        title = doc.add_heading('REPORTE DE ANÁLISIS LEGAL', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Información básica
        doc.add_heading('I. INFORMACIÓN GENERAL', level=1)
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ("Consulta:", report_data['query']),
            ("Fecha:", report_data['generation_date']),
            ("Hora:", report_data['generation_time']),
            ("Trace ID:", report_data['trace_id']),
            ("Método:", report_data['method']),
            ("Confianza:", f"{report_data['confidence']:.1%}")
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = str(value)
        
        # Base legal
        doc.add_heading('II. BASE LEGAL', level=1)
        if report_data['legal_basis']:
            for i, legal_item in enumerate(report_data['legal_basis'], 1):
                doc.add_heading(f'{i}. {legal_item["norm"]}', level=2)
                doc.add_paragraph(f"Artículo/Sección: {legal_item['article']}")
                doc.add_paragraph(f"Contenido: {legal_item['content']}")
                doc.add_paragraph(f"Relevancia: {legal_item['relevance']:.2f}")
        else:
            doc.add_paragraph("No se identificó base legal específica.")
        
        # Análisis
        doc.add_heading('III. ANÁLISIS', level=1)
        doc.add_paragraph(report_data['response'])
        
        # Referencias normativas
        if report_data['normative_references']:
            doc.add_heading('IV. REFERENCIAS NORMATIVAS', level=1)
            for ref in report_data['normative_references']:
                p = doc.add_paragraph()
                p.add_run(f"• {ref}")
        
        # Conclusiones
        if report_data['conclusions']:
            doc.add_heading('V. CONCLUSIONES', level=1)
            for i, conclusion in enumerate(report_data['conclusions'], 1):
                doc.add_paragraph(f"{i}. {conclusion}")
        
        # Metadatos técnicos
        doc.add_heading('VI. INFORMACIÓN TÉCNICA', level=1)
        tech_para = doc.add_paragraph()
        tech_para.add_run(f"Sistema: {report_data['system_version']}\n")
        tech_para.add_run(f"Documentos consultados: {report_data['documents_found']}\n")
        tech_para.add_run(f"Agente utilizado: {report_data['agent_used']}\n")
        
        # Guardar archivo
        filename = f"reporte_legal_{report_data['trace_id']}.docx"
        output_file = self.output_path / filename
        doc.save(str(output_file))
        
        logger.info(f"✅ Reporte Word generado: {output_file}")
        return output_file
    
    def _generate_pdf_report(self, report_data: Dict[str, Any], template_name: str) -> Path:
        """Generar reporte en formato PDF"""
        if not PDF_AVAILABLE:
            raise ImportError("WeasyPrint no está disponible")
        
        # Generar HTML primero
        html_content = self._generate_html_content(report_data, template_name)
        
        # Convertir a PDF
        filename = f"reporte_legal_{report_data['trace_id']}.pdf"
        output_file = self.output_path / filename
        
        weasyprint.HTML(string=html_content).write_pdf(str(output_file))
        
        logger.info(f"✅ Reporte PDF generado: {output_file}")
        return output_file
    
    def _generate_html_report(self, report_data: Dict[str, Any], template_name: str) -> Path:
        """Generar reporte en formato HTML"""
        html_content = self._generate_html_content(report_data, template_name)
        
        filename = f"reporte_legal_{report_data['trace_id']}.html"
        output_file = self.output_path / filename
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        logger.info(f"✅ Reporte HTML generado: {output_file}")
        return output_file
    
    def _generate_html_content(self, report_data: Dict[str, Any], template_name: str) -> str:
        """Generar contenido HTML usando Jinja2"""
        try:
            template = self.jinja_env.get_template(f"{template_name}.html")
            return template.render(**report_data)
        except Exception:
            # Usar template básico si no existe el específico
            return self._generate_basic_html(report_data)
    
    def _generate_basic_html(self, report_data: Dict[str, Any]) -> str:
        """Generar HTML básico sin template"""
        html_template = """
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Reporte Legal - {{ trace_id }}</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 20px; line-height: 1.6; }
        .header { text-align: center; border-bottom: 2px solid #333; padding-bottom: 10px; margin-bottom: 20px; }
        .section { margin-bottom: 30px; }
        .section h2 { color: #2c5aa0; border-bottom: 1px solid #ccc; padding-bottom: 5px; }
        .info-table { width: 100%; border-collapse: collapse; margin-bottom: 20px; }
        .info-table th, .info-table td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        .info-table th { background-color: #f2f2f2; }
        .legal-basis { background-color: #f9f9f9; padding: 15px; border-left: 4px solid #2c5aa0; margin: 10px 0; }
        .conclusions { background-color: #e8f5e8; padding: 15px; border-radius: 5px; }
        .footer { margin-top: 40px; padding-top: 20px; border-top: 1px solid #ccc; font-size: 0.9em; color: #666; }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ institution }}</h1>
        <h2>REPORTE DE ANÁLISIS LEGAL</h2>
        <p>Fecha: {{ generation_date }} - Hora: {{ generation_time }}</p>
    </div>
    
    <div class="section">
        <h2>I. INFORMACIÓN GENERAL</h2>
        <table class="info-table">
            <tr><th>Consulta</th><td>{{ query }}</td></tr>
            <tr><th>Trace ID</th><td>{{ trace_id }}</td></tr>
            <tr><th>Método</th><td>{{ method }}</td></tr>
            <tr><th>Agente</th><td>{{ agent_used }}</td></tr>
            <tr><th>Confianza</th><td>{{ "%.1f%%"|format(confidence * 100) }}</td></tr>
            <tr><th>Documentos</th><td>{{ documents_found }}</td></tr>
        </table>
    </div>
    
    {% if legal_basis %}
    <div class="section">
        <h2>II. BASE LEGAL</h2>
        {% for legal in legal_basis %}
        <div class="legal-basis">
            <h3>{{ legal.norm }}</h3>
            <p><strong>Artículo:</strong> {{ legal.article }}</p>
            <p><strong>Contenido:</strong> {{ legal.content }}</p>
            <p><strong>Relevancia:</strong> {{ "%.2f"|format(legal.relevance) }}</p>
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    <div class="section">
        <h2>III. ANÁLISIS</h2>
        <div>{{ response|replace('\n', '<br>')|safe }}</div>
    </div>
    
    {% if normative_references %}
    <div class="section">
        <h2>IV. REFERENCIAS NORMATIVAS</h2>
        <ul>
        {% for ref in normative_references %}
            <li>{{ ref }}</li>
        {% endfor %}
        </ul>
    </div>
    {% endif %}
    
    {% if conclusions %}
    <div class="section">
        <h2>V. CONCLUSIONES</h2>
        <div class="conclusions">
        {% for conclusion in conclusions %}
            <p>{{ loop.index }}. {{ conclusion }}</p>
        {% endfor %}
        </div>
    </div>
    {% endif %}
    
    <div class="footer">
        <p>Generado por {{ system_version }} el {{ generation_date }} a las {{ generation_time }}</p>
        <p>Este reporte ha sido generado automáticamente y debe ser validado por personal técnico competente.</p>
    </div>
</body>
</html>
        """
        
        template = Template(html_template)
        return template.render(**report_data)
    
    def _setup_word_styles(self, doc):
        """Configurar estilos para documento Word"""
        try:
            # Estilo para encabezados
            styles = doc.styles
            
            # Crear estilo personalizado si no existe
            try:
                if DOCX_AVAILABLE:
                    heading_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
                    heading_style.font.name = 'Arial'
                    heading_style.font.size = Pt(14)
                    heading_style.font.bold = True
            except Exception:
                pass  # El estilo ya existe o DOCX no disponible
                
        except Exception as e:
            logger.warning(f"No se pudieron configurar estilos Word: {e}")
    
    def _load_available_templates(self) -> List[ReportTemplate]:
        """Cargar templates disponibles"""
        templates = [
            ReportTemplate(
                name="legal_analysis",
                title="Análisis Legal",
                description="Reporte de análisis legal con base normativa",
                sections=["info", "legal_basis", "analysis", "conclusions"],
                format_type=ReportFormat.WORD,
                template_file="legal_analysis.html"
            ),
            ReportTemplate(
                name="audit_report",
                title="Reporte de Auditoría",
                description="Reporte de auditoría normativa",
                sections=["info", "findings", "recommendations"],
                format_type=ReportFormat.PDF,
                template_file="audit_report.html"
            )
        ]
        return templates
    
    def _ensure_basic_templates(self):
        """Crear templates básicos si no existen"""
        basic_template_path = self.templates_path / "legal_analysis.html"
        
        if not basic_template_path.exists():
            # El template HTML básico se genera dinámicamente en _generate_basic_html
            logger.info("Usando template HTML dinámico")
    
    def list_available_templates(self) -> List[Dict[str, Any]]:
        """Listar templates disponibles"""
        return [
            {
                "name": template.name,
                "title": template.title,
                "description": template.description,
                "sections": template.sections,
                "format": template.format_type.value
            }
            for template in self.available_templates
        ]
    
    def get_generation_capabilities(self) -> Dict[str, Any]:
        """Obtener capacidades de generación"""
        return {
            "formats_available": {
                "word": DOCX_AVAILABLE,
                "pdf": PDF_AVAILABLE,
                "html": True
            },
            "templates_loaded": len(self.available_templates),
            "output_path": str(self.output_path),
            "templates_path": str(self.templates_path),
            "dependencies": {
                "jinja2": True,
                "python_docx": DOCX_AVAILABLE,
                "weasyprint": PDF_AVAILABLE
            }
        }

# Funciones de utilidad
def generate_quick_report(
    query: str,
    analysis_result: Dict[str, Any],
    format_type: ReportFormat = ReportFormat.HTML,
    output_dir: Optional[Path] = None
) -> Path:
    """Función helper para generar reporte rápido"""
    generator = NormativeReportGenerator(output_path=output_dir)
    return generator.generate_legal_report(query, analysis_result, format_type=format_type)

# Instancia global
global_report_generator = NormativeReportGenerator()